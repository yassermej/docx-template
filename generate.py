from itertools import count
import os
import json
import re
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FILES_DIR = os.path.join(BASE_DIR, 'files')

REPLACE_PATTERN = re.compile(r'{{ (.*?) }}')

def get_nested_value(data, key_list):
    current_key = key_list.pop()
    if len(key_list) == 0:
        print("current_key: ", current_key, data[current_key])
        if data is None:
            return None
        return data[current_key]
    else:
        if current_key in data:
            return get_nested_value(data[current_key], key_list)

def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def insert_paragraph_after(paragraph, text=None, style=None, is_bold=None, color=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = None
    if text:
        run = new_para.add_run(text)
    if style is not None:
        new_para.style = style
    if is_bold is not None:
        run.font.bold = is_bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
        print("~~~~~~: ", color)

    return new_para

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


if __name__ == "__main__":
    TEMPLATE_PATH = os.path.join(FILES_DIR, 'template.docx')
    OUTPUT_PATH = os.path.join(FILES_DIR, 'output.docx')
    INPUT_PATH = os.path.join(FILES_DIR, 'data_text.txt')

    REMOVE_LIST = [
        '[ as bullets ]',
        'database.db.name',
        'database.db.link',
        'Search Strategy',
        'Search Terms SoTA',
        'Search Terms S&P',
        '{{ for item in database.sota_terms }}',
        '{{ for item in database.sp_terms }}',
        '[ as bullets ]',
    ]

    MATCH_REMOVE_LIST = [
        'item in database.sota_terms',
        'item in database.sp_terms',
        'database.db.name',
        'database.db.link'
    ]

    # Opening JSON file
    f = open(INPUT_PATH, encoding='utf-8')
    data = json.load(f, strict=False)
    replacable_keys = data.keys()

    doc = Document(TEMPLATE_PATH)

    graph_index = 0
    sci_database_index = 0
    for graph in doc.paragraphs:
        for run in graph.runs:
            if run.text.strip() in REMOVE_LIST:
                run.text = ''
            if 'Scientific Databases' == run.text:
                sci_database_index = graph_index
            else:
                matches = re.findall(REPLACE_PATTERN, run.text)
                for match in matches:
                    # Insert data
                    if match == 'date':
                        run.text = datetime.today().strftime("%b %d, %Y")

                    elif 'inclusion_criteria' in  match:
                        temp = ''
                        index = 1
                        for entry in data['inclusion_criteria']:
                            if index == 1:
                                graph.style.name = 'List Bullet'
                                list_number(doc, graph, level=0, num=False)
                                run.text = entry
                            else:
                                criteria_graph = insert_paragraph_after(graph, entry, 'List Bullet')
                                list_number(doc, criteria_graph, level=0, num=False)
                            index = index + 1
                    
                    elif 'exclusion_criteria' in  match:
                        temp = ''
                        index = 1
                        for entry in data['exclusion_criteria']:
                            if index == 1:
                                graph.style.name = 'List Bullet'
                                list_number(doc, graph, level=0, num=False)
                                run.text = entry
                            else:
                                exclusion_criteria_graph = insert_paragraph_after(graph, entry, 'List Bullet')
                                list_number(doc, exclusion_criteria_graph, level=0, num=False)
                            index = index + 1

                    elif 'database in sci_databases' in  match:
                        temp = ''
                        index = 1
                        for database in data['sci_databases']:
                            if index == 1:
                                run.text = ''
                            else:
                                db_name_graph = insert_paragraph_after(graph, database['db']['name'], None, True, '548DD4')
                                db_link_graph = insert_paragraph_after(db_name_graph, database['db']['link'])

                                empty_before_search_startegy = insert_paragraph_after(db_link_graph, '')

                                search_startegy_graph = insert_paragraph_after(empty_before_search_startegy, 'Search Strategy', None, True, '548DD4')



                                empty_after_search_startegy = insert_paragraph_after(search_startegy_graph, '')

                                search_so_term_graph = insert_paragraph_after(empty_after_search_startegy, 'Search Terms SoTA', None, True, '548DD4')

                                empty_after_search_so = insert_paragraph_after(search_so_term_graph, '')

                                sota_count = 1
                                sota_terms_graph = None
                                for term in database['sota_terms']:
                                    if sota_count == 1:
                                        sota_terms_graph = insert_paragraph_after(empty_after_search_so, term, 'List Bullet')
                                        list_number(doc, sota_terms_graph, level=0, num=False)
                                    else:
                                        sota_terms_graph = insert_paragraph_after(sota_terms_graph, term, 'List Bullet')
                                        list_number(doc, sota_terms_graph, level=0, num=False)
                                    sota_count = sota_count + 1


                                empty_before_new_graph = insert_paragraph_after(sota_terms_graph, '')
                                new_graph = insert_paragraph_after(empty_before_new_graph, 'Search Terms S&P', None, True, '548DD4')
                                empty_after_new_graph = insert_paragraph_after(new_graph, '')
                                new_graph.style.font.bold = True
                                for sp_term in database['sp_terms']:
                                    new_graph = insert_paragraph_after(empty_after_new_graph, sp_term, 'List Bullet')
                                    list_number(doc, new_graph, level=0, num=False)

                            index = index + 1

                    elif 'suitability_table_caption' in match:
                        run.text = data['suitability_table']['table_caption']

                    elif 'suitability_criteria_table_caption' in match:
                        run.text = data['suitability_criteria_table']['table_caption']

                    elif match in MATCH_REMOVE_LIST:
                        run.text = ''

                    # Replace variables
                    elif '.' in match:
                        key_list = match.split('.')
                        key_list.reverse()
                        nested_value = get_nested_value(data, key_list)
                        if nested_value != None:
                            run.text = nested_value
                    elif match in replacable_keys:
                        temp = run.text
                        variable_name = match.replace('{{', '').replace('}}', '').strip()
                        temp = temp.replace('{{ ' + match + ' }}', data[variable_name])
                        run.text = temp
        graph_index = graph_index + 1


    for table in doc.tables:
        if table.cell(0, 0).text.strip() == 'Criteria':
            suitability_table_headers = data['suitability_table']['headers']
            suitability_table_rows = data['suitability_table']['rows']
            count_rows = len(suitability_table_rows)
            row_id = 0
            for row in table.rows:
                if row_id == 0:
                    pass
                else:
                    if row_id <= count_rows:
                        row.cells[0].text = suitability_table_rows[row_id - 1][suitability_table_headers[0]]
                        row.cells[1].text = suitability_table_rows[row_id - 1][suitability_table_headers[1]]
                    else:
                        remove_row(table, row)
                row_id = row_id + 1
            if (row_id < count_rows):
                row = table.add_row()
                row.cells[0].text = suitability_table_rows[row_id - 1][suitability_table_headers[0]]
                row.cells[1].text = suitability_table_rows[row_id - 1][suitability_table_headers[1]]
                row_id = row_id + 1
            
            table_element = table._element
            parent = table_element.getparent()

        elif table.cell(0, 0).text.strip() == 'Suitability Criteria':
            suitability_criteria_table_headers = data['suitability_criteria_table']['headers']
            suitability_criteria_table_rows = data['suitability_criteria_table']['rows']
            count_rows = len(suitability_criteria_table_rows)
            row_id = 0
            for row in table.rows:
                if row_id == 0:
                    pass
                else:
                    if row_id <= count_rows:
                        row.cells[0].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[0]]
                        row.cells[1].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[1]]
                        row.cells[2].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[2]]
                    else:
                        remove_row(table, row)
                row_id = row_id + 1
            if (row_id < count_rows):
                row = table.add_row()
                row.cells[0].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[0]]
                row.cells[1].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[1]]
                row.cells[2].text = suitability_criteria_table_rows[row_id - 1][suitability_criteria_table_headers[2]]
                row_id = row_id + 1
    # Now save the document
    doc.save('output.docx')